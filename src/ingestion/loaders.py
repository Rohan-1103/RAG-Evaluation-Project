"""
src/ingestion/loaders.py

Concrete loader implementations for PDF, TXT/Markdown, and HTML formats.

Each loader:
  1. Implements BaseLoader (supported_extensions + _load_file)
  2. Populates DocumentMetadata as completely as the format allows
  3. Raises LoaderError on unrecoverable parse failures
  4. Never touches chunking, embedding, or storage

Loader responsibilities by format:
  PDFLoader   — pypdf page-by-page extraction, metadata from PDF info dict
  TXTLoader   — plain text + markdown, whole-file as single Document
  HTMLLoader  — beautifulsoup4 text extraction, strips scripts/styles

Format detection is by file extension only — content sniffing is
deliberately excluded. If a file has the wrong extension, rename it.
Attempting to detect format from content bytes adds complexity with
no meaningful benefit for a controlled ingestion pipeline.

Adding a new format:
  1. Create a new class inheriting BaseLoader
  2. Implement supported_extensions and _load_file
  3. Register it in LoaderRegistry in src/ingestion/pipeline.py
  Zero changes to any existing loader or pipeline orchestration code.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from loguru import logger

from config.settings import IngestionConfig
from src.ingestion.base import (
    BaseLoader,
    Document,
    DocumentMetadata,
    LoaderError,
)


# ===========================================================================
# PDF LOADER
# ===========================================================================


class PDFLoader(BaseLoader):
    """
    Load PDF files using pypdf.

    Produces one Document per page. Page-level granularity is
    intentional — it preserves page number metadata which is
    critical for vectorless RAG citations and eval drilldown.

    Merging pages before chunking loses page boundary information
    permanently. Chunking page-level documents preserves it.

    PDF metadata extraction:
      pypdf exposes /Title, /Author, /CreationDate from the PDF
      info dict. These are populated when present and None otherwise.
      Encrypted PDFs are attempted with an empty password — if that
      fails, LoaderError is raised with a clear message.
    """

    def __init__(self, config: IngestionConfig) -> None:
        self._config = config
        self._max_pages = config.pdf_max_pages   # 0 = no limit

    @property
    def supported_extensions(self) -> frozenset[str]:
        return frozenset({"pdf"})

    def _load_file(self, path: Path) -> list[Document]:
        """
        Load all pages from a PDF file.

        Returns one Document per non-empty page.
        Page numbers are 1-based to match human-readable PDF pages.
        """
        try:
            from pypdf import PdfReader
            from pypdf.errors import PdfReadError
        except ImportError as exc:
            raise LoaderError(
                loader="PDFLoader",
                path=path,
                reason=(
                    "pypdf is not installed. Run: poetry add pypdf"
                ),
                original_exception=exc,
            ) from exc

        try:
            reader = PdfReader(str(path))
        except Exception as exc:
            raise LoaderError(
                loader="PDFLoader",
                path=path,
                reason=f"pypdf failed to open file: {exc}",
                original_exception=exc,
            ) from exc

        # Handle encrypted PDFs
        if reader.is_encrypted:
            try:
                result = reader.decrypt("")
                if result == 0:
                    raise LoaderError(
                        loader="PDFLoader",
                        path=path,
                        reason=(
                            "PDF is encrypted and could not be decrypted "
                            "with an empty password. Provide a decrypted "
                            "copy of this file."
                        ),
                    )
            except LoaderError:
                raise
            except Exception as exc:
                raise LoaderError(
                    loader="PDFLoader",
                    path=path,
                    reason=f"Failed to decrypt PDF: {exc}",
                    original_exception=exc,
                ) from exc

        total_pages = len(reader.pages)
        file_size = path.stat().st_size

        # Extract PDF-level metadata from info dict
        pdf_info = reader.metadata or {}
        pdf_title = self._safe_pdf_meta(pdf_info, "/Title")
        pdf_author = self._safe_pdf_meta(pdf_info, "/Author")
        pdf_created = self._safe_pdf_meta(pdf_info, "/CreationDate")

        # Determine page range
        max_pages = (
            self._max_pages
            if self._max_pages > 0
            else total_pages
        )
        pages_to_load = min(total_pages, max_pages)

        if self._max_pages > 0 and total_pages > self._max_pages:
            logger.warning(
                f"PDFLoader: '{path.name}' has {total_pages} pages but "
                f"pdf_max_pages={self._max_pages}. "
                f"Loading first {pages_to_load} pages only."
            )

        documents: list[Document] = []

        for page_idx in range(pages_to_load):
            page_number = page_idx + 1   # 1-based

            try:
                page = reader.pages[page_idx]
                text = page.extract_text() or ""
            except Exception as exc:
                logger.warning(
                    f"PDFLoader: Failed to extract text from page "
                    f"{page_number} of '{path.name}': {exc}. Skipping page."
                )
                continue

            # Clean extracted text
            text = self._clean_pdf_text(text)

            if not text.strip():
                logger.debug(
                    f"PDFLoader: Page {page_number} of '{path.name}' "
                    f"is empty after extraction. Skipping."
                )
                continue

            metadata = DocumentMetadata(
                source_file=path.name,
                source_path=str(path),
                file_type="pdf",
                file_size_bytes=file_size,
                total_pages=total_pages,
                page_number=page_number,
                title=pdf_title,
                author=pdf_author,
                created_at=pdf_created,
            )

            documents.append(
                Document(
                    page_content=text,
                    metadata=metadata,
                )
            )

        logger.info(
            f"PDFLoader: Loaded {len(documents)} pages "
            f"from '{path.name}' "
            f"(total_pages={total_pages})."
        )

        return documents

    def _clean_pdf_text(self, text: str) -> str:
        """
        Clean raw pypdf text extraction output.

        pypdf sometimes produces:
          - Excessive whitespace between words (ligature splitting)
          - Hyphenated line breaks ("informa-\ntion")
          - Repeated newlines from column layouts

        We apply conservative cleaning — aggressive cleaning risks
        corrupting technical content (code, tables, equations).
        """
        # Rejoin hyphenated line breaks ("informa-\ntion" → "information")
        text = re.sub(r"-\n(\w)", r"\1", text)

        # Collapse 3+ consecutive newlines to 2 (preserve paragraph breaks)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Replace non-breaking spaces with regular spaces
        text = text.replace("\xa0", " ")

        # Collapse multiple spaces to single space
        text = re.sub(r" {2,}", " ", text)

        return text.strip()

    @staticmethod
    def _safe_pdf_meta(info: Any, key: str) -> str | None:
        """
        Safely extract a value from pypdf's metadata dict.

        pypdf metadata values can be:
          - str (normal case)
          - pypdf.generic.TextStringObject (needs str())
          - None (key not present)
          - Raises on access for malformed PDFs

        Returns None for any non-string or missing value.
        """
        try:
            val = info.get(key)
            if val is None:
                return None
            return str(val).strip() or None
        except Exception:
            return None


# ===========================================================================
# TXT / MARKDOWN LOADER
# ===========================================================================


class TXTLoader(BaseLoader):
    """
    Load plain text and Markdown files.

    Unlike PDFLoader (one Document per page), TXTLoader loads the
    entire file as a single Document. Text files have no inherent
    page structure — chunking happens downstream in RecursiveChunker.

    Encoding handling:
      Tries UTF-8 first (correct for 99% of modern text files).
      Falls back to latin-1 which never fails (every byte is valid).
      The encoding used is recorded in DocumentMetadata.extra so
      downstream modules can flag latin-1 fallbacks for review.

    Markdown support:
      .md and .markdown files are loaded identically to .txt.
      No markdown parsing (bold, links, headers) is performed —
      the raw markdown syntax is preserved. The chunker splits on
      \n\n which naturally respects markdown section boundaries.
    """

    def __init__(self, config: IngestionConfig) -> None:
        self._config = config

    @property
    def supported_extensions(self) -> frozenset[str]:
        return frozenset({"txt", "text", "md", "markdown"})

    def _load_file(self, path: Path) -> list[Document]:
        """
        Load entire text file as a single Document.

        Returns a list with exactly one Document (or zero if file
        is entirely whitespace after loading).
        """
        file_size = path.stat().st_size

        if file_size == 0:
            logger.warning(
                f"TXTLoader: '{path.name}' is an empty file (0 bytes). "
                f"Returning no documents."
            )
            return []

        # Try UTF-8 first, fall back to latin-1
        text, encoding_used = self._read_with_fallback(path)

        if encoding_used == "latin-1":
            logger.warning(
                f"TXTLoader: '{path.name}' could not be decoded as UTF-8. "
                f"Loaded with latin-1 fallback. "
                f"Consider converting the file to UTF-8."
            )

        # Clean text
        text = self._clean_text(text)

        if not text.strip():
            logger.warning(
                f"TXTLoader: '{path.name}' contains only whitespace "
                f"after cleaning. Returning no documents."
            )
            return []

        ext = path.suffix.lstrip(".").lower()

        metadata = DocumentMetadata(
            source_file=path.name,
            source_path=str(path),
            file_type=ext,
            file_size_bytes=file_size,
            total_pages=None,     # Text files have no pages
            page_number=None,
            title=self._infer_title(text, path),
            author=None,
            created_at=None,
            extra={"encoding": encoding_used},
        )

        logger.info(
            f"TXTLoader: Loaded '{path.name}' "
            f"({len(text)} chars, encoding={encoding_used})."
        )

        return [Document(page_content=text, metadata=metadata)]

    def _read_with_fallback(
        self,
        path: Path,
    ) -> tuple[str, str]:
        """
        Read file content, trying UTF-8 then latin-1.

        Returns (content, encoding_used).
        """
        try:
            content = path.read_text(encoding="utf-8")
            return content, "utf-8"
        except UnicodeDecodeError:
            pass
        except Exception as exc:
            raise LoaderError(
                loader="TXTLoader",
                path=path,
                reason=f"Failed to read file: {exc}",
                original_exception=exc,
            ) from exc

        try:
            content = path.read_text(encoding="latin-1")
            return content, "latin-1"
        except Exception as exc:
            raise LoaderError(
                loader="TXTLoader",
                path=path,
                reason=(
                    f"Failed to read file with both UTF-8 and latin-1 "
                    f"encodings: {exc}"
                ),
                original_exception=exc,
            ) from exc

    def _clean_text(self, text: str) -> str:
        """
        Conservative text cleaning for plain text files.

        Preserves intentional formatting (code blocks, lists, tables)
        while removing noise (null bytes, excessive blank lines).
        """
        # Remove null bytes (sometimes present in Windows-created files)
        text = text.replace("\x00", "")

        # Normalise Windows line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Collapse 4+ consecutive newlines to 2
        text = re.sub(r"\n{4,}", "\n\n\n", text)

        return text.strip()

    @staticmethod
    def _infer_title(text: str, path: Path) -> str | None:
        """
        Attempt to infer document title from content or filename.

        For Markdown: first # heading is the title.
        For plain text: first non-empty line if it looks like a title
                        (short, no sentence-ending punctuation).
        Falls back to filename stem.
        """
        lines = text.strip().splitlines()
        if not lines:
            return path.stem

        first_line = lines[0].strip()

        # Markdown heading
        if first_line.startswith("#"):
            return first_line.lstrip("#").strip() or path.stem

        # Short first line without sentence punctuation = likely a title
        if (
            len(first_line) <= 120
            and not first_line.endswith((".", "?", "!"))
            and first_line
        ):
            return first_line

        return path.stem


# ===========================================================================
# HTML LOADER
# ===========================================================================


class HTMLLoader(BaseLoader):
    """
    Load HTML files using BeautifulSoup4.

    Extracts visible text content, stripping:
      - <script> and <style> tags and their content
      - HTML comments
      - Navigation, header, footer elements (common boilerplate)

    Metadata extraction:
      <title> tag → DocumentMetadata.title
      <meta name="author"> → DocumentMetadata.author
      <meta name="description"> → stored in extra

    Why not use LangChain's BSHTMLLoader:
      LangChain's loader is a thin wrapper that doesn't strip
      scripts/styles or extract metadata. We need cleaner extraction
      and structured metadata for the evaluation pipeline.
    """

    # HTML tags whose entire content (tag + children) should be removed
    _REMOVE_TAGS: frozenset[str] = frozenset({
        "script",
        "style",
        "noscript",
        "iframe",
        "svg",
        "nav",
        "header",
        "footer",
        "aside",
        "form",
        "button",
        "input",
        "select",
        "textarea",
    })

    # Tags treated as block elements — get a newline before/after
    _BLOCK_TAGS: frozenset[str] = frozenset({
        "p", "div", "section", "article", "main",
        "h1", "h2", "h3", "h4", "h5", "h6",
        "li", "td", "th", "tr",
        "blockquote", "pre", "code",
        "br", "hr",
    })

    def __init__(self, config: IngestionConfig) -> None:
        self._config = config

    @property
    def supported_extensions(self) -> frozenset[str]:
        return frozenset({"html", "htm", "xhtml"})

    def _load_file(self, path: Path) -> list[Document]:
        """
        Parse HTML and extract clean text as a single Document.
        """
        try:
            from bs4 import BeautifulSoup, Comment
        except ImportError as exc:
            raise LoaderError(
                loader="HTMLLoader",
                path=path,
                reason=(
                    "beautifulsoup4 is not installed. "
                    "Run: poetry add beautifulsoup4 lxml"
                ),
                original_exception=exc,
            ) from exc

        file_size = path.stat().st_size

        # Read raw HTML
        try:
            raw_html = path.read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            raise LoaderError(
                loader="HTMLLoader",
                path=path,
                reason=f"Failed to read HTML file: {exc}",
                original_exception=exc,
            ) from exc

        # Parse with lxml for speed; fall back to html.parser
        try:
            soup = BeautifulSoup(raw_html, "lxml")
        except Exception:
            try:
                soup = BeautifulSoup(raw_html, "html.parser")
            except Exception as exc:
                raise LoaderError(
                    loader="HTMLLoader",
                    path=path,
                    reason=f"BeautifulSoup failed to parse HTML: {exc}",
                    original_exception=exc,
                ) from exc

        # Extract metadata before modifying soup
        title = self._extract_title(soup, path)
        author = self._extract_meta(soup, "author")
        description = self._extract_meta(soup, "description")

        # Remove noise tags and HTML comments
        self._strip_noise(soup)

        # Extract clean text
        text = self._extract_text(soup)
        text = self._clean_html_text(text)

        if not text.strip():
            logger.warning(
                f"HTMLLoader: '{path.name}' produced no text "
                f"after extraction. Returning no documents."
            )
            return []

        extra: dict[str, Any] = {}
        if description:
            extra["meta_description"] = description

        metadata = DocumentMetadata(
            source_file=path.name,
            source_path=str(path),
            file_type="html",
            file_size_bytes=file_size,
            total_pages=None,
            page_number=None,
            title=title,
            author=author,
            created_at=None,
            extra=extra,
        )

        logger.info(
            f"HTMLLoader: Loaded '{path.name}' "
            f"({len(text)} chars, title='{title}')."
        )

        return [Document(page_content=text, metadata=metadata)]

    def _strip_noise(self, soup: Any) -> None:
        """
        Remove noise elements from parsed HTML in place.

        Removes: script, style, nav, footer, comments, and other
        non-content elements defined in _REMOVE_TAGS.
        """
        from bs4 import Comment

        # Remove tags defined in _REMOVE_TAGS
        for tag_name in self._REMOVE_TAGS:
            for tag in soup.find_all(tag_name):
                tag.decompose()

        # Remove HTML comments
        for comment in soup.find_all(
            string=lambda text: isinstance(text, Comment)
        ):
            comment.extract()

    def _extract_text(self, soup: Any) -> str:
        """
        Extract text from parsed soup with block element newlines.

        Inserts newlines before block elements so the output text
        preserves paragraph structure rather than concatenating
        all text into one long line.
        """
        # Insert newlines before block elements
        for tag in soup.find_all(self._BLOCK_TAGS):
            tag.insert_before("\n")
            tag.insert_after("\n")

        # Get all text
        text = soup.get_text(separator=" ", strip=False)
        return text

    def _clean_html_text(self, text: str) -> str:
        """
        Clean text extracted from HTML.

        HTML extraction produces:
          - Many consecutive blank lines (from block element newlines)
          - Trailing/leading spaces on lines
          - Lines that are only whitespace
        """
        lines = text.splitlines()
        cleaned_lines: list[str] = []

        for line in lines:
            stripped = line.strip()
            cleaned_lines.append(stripped)

        text = "\n".join(cleaned_lines)

        # Collapse 3+ consecutive newlines to 2
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove lines that are only punctuation/symbols (nav artifacts)
        lines = text.splitlines()
        content_lines = [
            line for line in lines
            if not re.match(r"^[^\w\s]*$", line) or not line.strip()
        ]
        text = "\n".join(content_lines)

        return text.strip()

    @staticmethod
    def _extract_title(soup: Any, path: Path) -> str | None:
        """Extract page title from <title> tag or first <h1>."""
        # Try <title> tag
        title_tag = soup.find("title")
        if title_tag and title_tag.get_text(strip=True):
            return title_tag.get_text(strip=True)

        # Try first <h1>
        h1 = soup.find("h1")
        if h1 and h1.get_text(strip=True):
            return h1.get_text(strip=True)

        return path.stem

    @staticmethod
    def _extract_meta(soup: Any, name: str) -> str | None:
        """
        Extract <meta name="..."> content attribute.

        Handles both name= and property= meta tags (Open Graph).
        """
        # Standard meta tag: <meta name="author" content="...">
        tag = soup.find("meta", attrs={"name": name})
        if tag and tag.get("content"):
            return str(tag["content"]).strip() or None

        # Open Graph / property meta: <meta property="og:author" ...>
        tag = soup.find("meta", attrs={"property": f"og:{name}"})
        if tag and tag.get("content"):
            return str(tag["content"]).strip() or None

        return None


# ===========================================================================
# DOCX LOADER
# ===========================================================================


class DOCXLoader(BaseLoader):
    """
    Load Microsoft Word (.docx) files using python-docx.

    Extracts paragraph text preserving heading hierarchy.
    Each paragraph becomes a text block; the full document is
    returned as a single Document (chunking is downstream).

    Heading detection:
      python-docx exposes paragraph.style.name which includes
      "Heading 1", "Heading 2" etc. Headings are prefixed with
      markdown-style # markers so the chunker can split on them.
    """

    def __init__(self, config: IngestionConfig) -> None:
        self._config = config

    @property
    def supported_extensions(self) -> frozenset[str]:
        return frozenset({"docx"})

    def _load_file(self, path: Path) -> list[Document]:
        """Load DOCX as a single Document with heading markers."""
        try:
            import docx
        except ImportError as exc:
            raise LoaderError(
                loader="DOCXLoader",
                path=path,
                reason=(
                    "python-docx is not installed. "
                    "Run: poetry add python-docx"
                ),
                original_exception=exc,
            ) from exc

        try:
            doc = docx.Document(str(path))
        except Exception as exc:
            raise LoaderError(
                loader="DOCXLoader",
                path=path,
                reason=f"python-docx failed to open file: {exc}",
                original_exception=exc,
            ) from exc

        file_size = path.stat().st_size
        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # Convert Word headings to markdown-style prefixes
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                    prefix = "#" * min(level, 6)
                    lines.append(f"{prefix} {text}")
                except (ValueError, IndexError):
                    lines.append(text)
            else:
                lines.append(text)

        full_text = "\n\n".join(lines)

        if not full_text.strip():
            logger.warning(
                f"DOCXLoader: '{path.name}' produced no text. "
                f"Returning no documents."
            )
            return []

        # Extract core properties
        props = doc.core_properties
        title = getattr(props, "title", None) or path.stem
        author = getattr(props, "author", None)
        created = None
        if getattr(props, "created", None):
            try:
                created = props.created.isoformat()
            except Exception:
                created = None

        metadata = DocumentMetadata(
            source_file=path.name,
            source_path=str(path),
            file_type="docx",
            file_size_bytes=file_size,
            total_pages=None,
            page_number=None,
            title=str(title) if title else None,
            author=str(author) if author else None,
            created_at=created,
        )

        logger.info(
            f"DOCXLoader: Loaded '{path.name}' "
            f"({len(full_text)} chars, "
            f"{len(doc.paragraphs)} paragraphs)."
        )

        return [Document(page_content=full_text, metadata=metadata)]


__all__ = [
    "PDFLoader",
    "TXTLoader",
    "HTMLLoader",
    "DOCXLoader",
]